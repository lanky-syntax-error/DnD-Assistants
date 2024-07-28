# GUBS
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import math
print("NEW NPC IN PROGRESS; PLEASE LEAVE ANY UNKNOWN VALUES BLANK")
doc = Document()

# INPUTS
name = input("NPC name: ")
race = input("NPC race: ")
age = input("NPC age: ")
alignment = input("NPC alignment: ")
ac = input("NPC armour class: ")
health = input("NPC HP: ")
speed = input("NPC speed (ft): ")
desc = input("Enter a short description of the NPC: ")

# BASE STATS
str = input("Enter NPC strength score: ")
str = int(str)
str_mod = math.floor((str - 10) / 2)
dex = input("Enter NPC dexterity score: ")
dex = int(dex)
dex_mod = math.floor((dex - 10) / 2)
con = input("Enter NPC constitution score: ")
con = int(con)
con_mod = math.floor((con - 10) / 2)
intel = input("Enter NPC intelligence score: ")
intel = int(intel)
intel_mod = math.floor((intel - 10) / 2)
wis = input("Enter NPC wisdom score: ")
wis = int(wis)
wis_mod = math.floor((wis - 10) / 2)
cha = input("Enter NPC charisma score: ")
cha = int(cha)
cha_mod = math.floor((cha - 10) / 2)

# ACTIONS
moves = input("Press enter to add a new combat action, or type 'DONE' to move on: ")
move_names = []
move_descriptions = []
move_checks = []

while moves != "DONE":
    move_name = input(f"Enter action name: ")
    move_description = input(f"Enter description of action: ")
    move_check = input("Enter the check/save required for this action: ")
    move_names.append(move_name)
    move_descriptions.append(move_description)
    move_checks.append(move_check)
    moves = input("Press enter to add a new combat action, or type 'DONE' to move on: ")

# BONUS ACTIONS
bonusmoves = input("Press enter to add a new bonus action, or type 'DONE' to move on: ")
bonusmove_names = []
bonusmove_descriptions = []
bonusmove_checks = []

while bonusmoves != "DONE":
    bonusmove_name = input(f"Enter bonus action name: ")
    bonusmove_description = input(f"Enter description of bonus action: ")
    bonusmove_check = input("Enter the check/save required for this bonus action: ")
    bonusmove_names.append(bonusmove_name)
    bonusmove_descriptions.append(bonusmove_description)
    bonusmove_checks.append(bonusmove_check)
    bonusmoves = input("Press enter to add a new bonus action, or type 'DONE' to move on: ")

# REACTIONS
reactions = input("Press enter to add a new reaction, or type 'DONE' to move on: ")
reaction_names = []
reaction_descriptions = []
reaction_checks = []

while reactions != "DONE":
    reaction_name = input(f"Enter reaction name: ")
    reaction_description = input(f"Enter description of reaction: ")
    reaction_check = input("Enter the check/save required for this reaction: ")
    reaction_names.append(reaction_name)
    reaction_descriptions.append(reaction_description)
    reaction_checks.append(reaction_check)
    reactions = input("Press enter to add a new reaction, or type 'DONE' to move on: ")

# RESISTANCES
resistance = input("Press enter to add a resistance for the NPC, or type 'DONE' to move on: ")
resistance_list = []
resistance_degree_list = []

while resistance != "DONE":
    resistance_nature = input("What is the NPC resistant to? ")
    resistance_degree = input(f"How resistant is the NPC to {resistance_nature}? Resistant/Immune? ")
    resistance_list.append(resistance_nature)
    resistance_degree_list.append(resistance_degree)
    resistance = input("Press enter to add a resistance for the NPC, or type 'DONE' to move on: ")

# WEAKNESSES   
weakness = input("Press enter to add a weakness for the NPC, or type 'DONE' to finish NPC creation: ")
weakness_list = []

while weakness != "DONE":
    weakness_name = input("Enter weakness: ")
    weakness_list.append(weakness_name)
    weakness = input("Press enter to add a weakness for the NPC, or type 'DONE' to finish NPC creation: ")


# PRINTING
doc.add_heading(f"{name}\n")

paragraph = doc.add_paragraph(f'{alignment} {race}\nAge: {age}\nSpeed: {speed}ft')
run = paragraph.runs[0]
run.bold = True

paragraph = doc.add_paragraph(f'AC: {ac}\nHEALTH: {health}')
run = paragraph.runs[0]
run.bold = True

paragraph = doc.add_paragraph(f'{desc}\n')
run = paragraph.runs[0]
run.italic = True

# BASE STATS
table = doc.add_table(rows=0, cols=6)
table.style = 'Table Grid'
row_cells = table.add_row().cells
row_cells[0].text = (f"STR: {str}({str_mod})")
row_cells[1].text = (f"DEX: {dex}({dex_mod})")
row_cells[2].text = (f"CON: {con}({con_mod})")
row_cells[3].text = (f"INT: {intel}({intel_mod})")
row_cells[4].text = (f"WIS: {wis}({wis_mod})")
row_cells[5].text = (f"CHA: {cha}({cha_mod})")

# ACTION LIST
if move_names:
    doc.add_heading(f"\nCombat actions", level=2)
table = doc.add_table(rows=0, cols=3)
table.style = 'Table Grid'

for move_name, move_description, move_check in zip(move_names, move_descriptions, move_checks):
    row_cells = table.add_row().cells
    row_cells[0].text = move_name
    row_cells[1].text = move_description
    row_cells[2].text = move_check

# BONUS ACTION LIST
if bonusmove_names:
    doc.add_heading("Bonus actions", level=2)
table = doc.add_table(rows=0, cols=3)
table.style = 'Table Grid'

for bonusmove_name, bonusmove_description, bonusmove_check in zip(bonusmove_names, bonusmove_descriptions, bonusmove_checks):
    row_cells = table.add_row().cells
    row_cells[0].text = bonusmove_name
    row_cells[1].text = bonusmove_description
    row_cells[2].text = bonusmove_check

# REACTION LIST
if reaction_names:
    doc.add_heading("Reactions", level=2)
table = doc.add_table(rows=0, cols=3)
table.style = 'Table Grid'

for reaction_name, reaction_description, reaction_check in zip(reaction_names, reaction_descriptions, reaction_checks):
    row_cells = table.add_row().cells
    row_cells[0].text = reaction_name
    row_cells[1].text = reaction_description
    row_cells[2].text = reaction_check

# RESISTANCE LIST
if resistance_list:
    doc.add_heading("Resistances/Immunities", level=2)
table = doc.add_table(rows=0, cols=2)
table.style = 'Table Grid'

for resistance_nature, resistance_degree in zip(resistance_list, resistance_degree_list):
    row_cells = table.add_row().cells
    row_cells[0].text = resistance_nature
    row_cells[1].text = resistance_degree

# WEAKNESS LIST
if weakness_list:
    doc.add_heading("Weaknesses", level=2)

for weakness_name in weakness_list:
   paragraph = doc.add_paragraph(f"{weakness_name}")

#SAVING
doc.save(f"C:\\Users\\camdi\\OneDrive\\Documents\\Coding projects\\DnD coding\\{name} stat sheet.docx")
print(f"{name}'s stat sheet saved sucessfully, can't wait for the players to murder them...")